
import nltk
nltk.download('stopwords')
import os
from docx import Document
from pyresparser import ResumeParser
import csv
import spacy
from spacy.tokens import DocBin
from tqdm import tqdm
import json
import pickle
import locale
import re


def getpreferredencoding(do_setlocale = True):
    return "UTF-8"
    locale.getpreferredencoding = getpreferredencoding


def get_spacy_doc(file,data):
    nlp = spacy.blank('en')
    db = DocBin()


    for text, annot in tqdm(data):
        doc = nlp.make_doc(text)
        annot = annot['entities']
        ents = []
        entity_indices = []

        for start,end, label in annot:
            skip_entity = False

            for idx in range(start, end):
                if idx in entity_indices:
                    skip_entity = True
                    break

            if skip_entity == True:
                continue

            entity_indices = entity_indices  + list(range(start, end))

            try:
                span = doc.char_span(start, end, label = label, alignment_mode = 'strict')
            except:
                continue

            if span is None:
                # err_data = str([start, end] + " "  + "\n")
                err_data = f"error occured at [{start}, {end}] : some error \n"
                file.write(err_data)
            if span is not None:
                ents.append(span)


        try:
            doc.ents = ents
            db.add(doc)
        except:
            pass

    return db


def trim_entity_spans(data: list) -> list:
    """Removes leading and trailing white spaces from entity spans.

    Args:
        data (list): The data to be cleaned in spaCy JSON format.

    Returns:
        list: The cleaned data.
    """
    # invalid_span_tokens = re.compile(r'\s')
    invalid_span_tokens = re.compile(r"^\s+|\s+$")

    cleaned_data = []
    for text, annotations in data:
        entities = annotations['entities']
        valid_entities = []
        for start, end, label in entities:
            valid_start = start
            valid_end = end
            while valid_start < len(text) and invalid_span_tokens.match(
                    text[valid_start]):
                valid_start += 1
            while valid_end > 1 and invalid_span_tokens.match(
                    text[valid_end - 1]):
                valid_end -= 1
            valid_entities.append([valid_start, valid_end, label])
        cleaned_data.append([text, {'entities': valid_entities}])

    return cleaned_data


def get_extracted_skill(path):
    ResumeParser = pickle.load(open(r"C:\Users\HP\Desktop\5TH YEAR PROJECT\application\resume_parsing_model.pkl"))
    ResumeParser(path).get_extracted_data()
    try:
        pass
        doc = Document()
        with open(path, 'r',encoding='utf-8') as file:
            doc.add_paragraph(file.read())
        doc.save("text.docx")
        data = ResumeParser('text.docx').get_extracted_data()
    
    except:
        data = ResumeParser(path).get_extracted_data()
        print(data)
        candidate['name'] = data['name'];
        candidate['skills'] = data['skills'];
        candidate['experience'] = data['experience']; 
        candidate['email'] = data['email']; 

        with open('cv_data.csv', mode = 'w') as csvfile:
            fieldnames = ["name","email", "skills","experience"]
            writer = csv.DictWriter(csvfile, fieldnames = fieldnames)
            writer.writeheader()
            writer.writerow(candidate)
        return candidate




def clean_entity_spans(data: list) -> list:
  invalid_span_tokens = re.compile(r'\s')

  cleaned_data = []

  for content in data:
      name = content['documentName']
      text = content['document']
      userinput = content['user_input']

      valid_entities = []

      for annotate_content in content['annotation']:
          start = annotate_content['start']
          end = annotate_content['end']
          label = annotate_content['label']
          text1 = annotate_content['text']

          valid_start = start
          valid_end = end

          while valid_start < len(text) and invalid_span_tokens.match(
                  text[valid_start]):
              valid_start += 1
          while valid_end > 1 and invalid_span_tokens.match(
                  text[valid_end - 1]):
              valid_end -= 1

          valid_entities.append({'start': valid_start, 'end': valid_end, 'label': label, 'text': text1, 'propertiesList': [], 'commentsList': []})
      cleaned_data.append({'documentName': name, 'document':text, 'annotation': valid_entities, 'user_input': userinput})

  return cleaned_data




candidate = {}
def get_extracted_skills(path):
    filed = path
    try:
        pass
        doc = Document()
        with open(path, 'r',encoding='utf-8') as file:
            doc.add_paragraph(file.read())
        doc.save("text.docx")
        data = ResumeParser('text.docx').get_extracted_data()
    
    except:
        data = ResumeParser(filed).get_extracted_data()
        print(data)
        candidate['name'] = data['name'];
        candidate['skills'] = data['skills'];
        candidate['experience'] = [data['experience']]; 
        candidate['email'] = data['email']; 
        candidate['total_experience'] = data['total_experience'];
        candidate['degree'] = data['degree'];

        # with open('cv_data.csv', mode = 'w') as csvfile:
        #     fieldnames = ["name","email", "skills","experience"]
        #     writer = csv.DictWriter(csvfile, fieldnames = fieldnames)
        #     writer.writeheader()
        #     writer.writerow(candidate)
        return candidate
    
# print(get_extracted_skills(r'C:\Users\HP\Desktop\5TH YEAR PROJECT\application\uploads\Sally_L_Ajevi.pdf'))